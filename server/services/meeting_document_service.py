"""
会议文档生成服务
"""
import io
from typing import Dict, Optional
from database import db
from models.transcript import Transcript
from utils.datetime_formatter import format_datetime_to_beijing

# 尝试导入docx库
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class MeetingDocumentService:
    """会议文档生成服务类"""
    
    def __init__(self, meeting_service):
        """
        初始化文档生成服务
        
        Args:
            meeting_service: MeetingService 实例，用于获取会议信息
        """
        self.meeting_service = meeting_service
    
    def generate_summary_document(self, meeting_id: str, user_id: Optional[int] = None) -> io.BytesIO:
        """
        生成会议总结Word文档
        
        Args:
            meeting_id: 会议ID
            user_id: 用户ID（用于权限检查）
        
        Returns:
            Word文档的字节流
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx 未安装，无法生成Word文档")
        
        # 获取会议信息
        meeting = self.meeting_service.get_meeting(meeting_id, user_id=user_id)
        if not meeting:
            raise ValueError(f"会议不存在: {meeting_id}")
        
        # 获取最新转写记录和摘要
        transcript_record = Transcript.query.filter_by(meeting_id=meeting_id).order_by(
            Transcript.created_at.desc()
        ).first()
        
        if not transcript_record or not transcript_record.summary_dict:
            raise ValueError("会议总结不存在，请先生成会议总结")
        
        summary_data = transcript_record.summary_dict
        
        # 创建Word文档
        doc = DocxDocument()
        self._setup_document_style(doc)
        self._add_title(doc, meeting.get('name', '会议总结'))
        self._add_basic_info(doc, meeting)
        self._add_summary_sections(doc, summary_data)
        
        # 将文档保存到内存
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    
    def _setup_document_style(self, doc: DocxDocument):
        """设置文档默认样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)
    
    def _add_title(self, doc: DocxDocument, title_text: str):
        """添加文档标题"""
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title.paragraph_format.space_after = Pt(12)
    
    def _add_basic_info(self, doc: DocxDocument, meeting: Dict):
        """添加会议基本信息"""
        doc.add_paragraph()  # 空行
        
        # 章节标题
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run('会议基本信息')
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = RGBColor(0, 51, 102)
        section_title.paragraph_format.space_before = Pt(6)
        section_title.paragraph_format.space_after = Pt(6)
        
        # 信息项
        info_items = []
        if meeting.get('subject'):
            info_items.append(f"学科：{meeting['subject']}")
        if meeting.get('grade'):
            info_items.append(f"年级：{meeting['grade']}")
        if meeting.get('lesson_type'):
            info_items.append(f"备课类型：{meeting['lesson_type']}")
        if meeting.get('created_at'):
            formatted_time = format_datetime_to_beijing(meeting['created_at'])
            info_items.append(f"创建时间：{formatted_time}")
        if meeting.get('teachers'):
            teacher_names = [t.get('name', '') for t in meeting['teachers'] if t.get('name')]
            if teacher_names:
                info_items.append(f"参会人员：{', '.join(teacher_names)}")
        
        for item in info_items:
            info_para = doc.add_paragraph()
            info_para.paragraph_format.left_indent = Inches(0.3)
            info_para.paragraph_format.space_after = Pt(3)
            info_run = info_para.add_run(item)
            info_run.font.size = Pt(11)
        
        doc.add_paragraph()  # 空行
    
    def _add_summary_sections(self, doc: DocxDocument, summary_data: Dict):
        """添加摘要各个章节"""
        self._add_full_summary(doc, summary_data)
        self._add_speaker_summary(doc, summary_data)
        self._add_qa_summary(doc, summary_data)
        self._add_key_points(doc, summary_data)
        self._add_mind_map(doc, summary_data)
    
    def _add_section_title(self, doc: DocxDocument, title: str):
        """添加章节标题"""
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run(title)
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = RGBColor(0, 51, 102)
        section_title.paragraph_format.space_before = Pt(6)
        section_title.paragraph_format.space_after = Pt(6)
        return section_title
    
    def _add_full_summary(self, doc: DocxDocument, summary_data: Dict):
        """添加全文摘要"""
        if summary_data.get('paragraph_summary') or summary_data.get('summary'):
            self._add_section_title(doc, '一、全文摘要')
            
            summary_text = summary_data.get('paragraph_summary') or summary_data.get('summary', '')
            summary_para = doc.add_paragraph(summary_text)
            summary_para.paragraph_format.left_indent = Inches(0.3)
            summary_para.paragraph_format.first_line_indent = Inches(0.3)
            summary_para.paragraph_format.line_spacing = 1.75
            summary_para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()  # 空行
    
    def _add_speaker_summary(self, doc: DocxDocument, summary_data: Dict):
        """添加发言总结"""
        conversational_summary = summary_data.get('conversational_summary', [])
        if not conversational_summary:
            return
        
        self._add_section_title(doc, '二、发言总结')
        
        for idx, speaker in enumerate(conversational_summary, 1):
            speaker_para = doc.add_paragraph()
            speaker_para.paragraph_format.left_indent = Inches(0.2)
            speaker_para.paragraph_format.space_before = Pt(6)
            speaker_name = speaker.get('SpeakerName') or speaker.get('SpeakerId') or f'发言人{idx}'
            speaker_run = speaker_para.add_run(f"{idx}. {speaker_name}")
            speaker_run.font.size = Pt(12)
            speaker_run.font.bold = True
            speaker_run.font.color.rgb = RGBColor(51, 102, 153)
            
            if speaker.get('Summary'):
                summary_para = doc.add_paragraph(speaker['Summary'])
                summary_para.paragraph_format.left_indent = Inches(0.5)
                summary_para.paragraph_format.first_line_indent = Inches(0.2)
                summary_para.paragraph_format.line_spacing = 1.6
                summary_para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # 空行
    
    def _add_qa_summary(self, doc: DocxDocument, summary_data: Dict):
        """添加问答回顾"""
        qa_summary = summary_data.get('questions_answering_summary', [])
        if not qa_summary:
            return
        
        self._add_section_title(doc, '三、问答回顾')
        
        for idx, qa in enumerate(qa_summary, 1):
            # 问题
            qa_para = doc.add_paragraph()
            qa_para.paragraph_format.left_indent = Inches(0.2)
            qa_para.paragraph_format.space_before = Pt(6)
            q_run = qa_para.add_run(f"{idx}. 问题：")
            q_run.font.size = Pt(12)
            q_run.font.bold = True
            q_run.font.color.rgb = RGBColor(153, 0, 0)
            
            if qa.get('Question'):
                question_para = doc.add_paragraph(qa['Question'])
                question_para.paragraph_format.left_indent = Inches(0.5)
                question_para.paragraph_format.first_line_indent = Inches(0.2)
                question_para.paragraph_format.line_spacing = 1.6
            
            # 回答
            answer_para = doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Inches(0.2)
            answer_para.paragraph_format.space_before = Pt(3)
            a_run = answer_para.add_run("   回答：")
            a_run.font.size = Pt(12)
            a_run.font.bold = True
            a_run.font.color.rgb = RGBColor(0, 102, 51)
            
            if qa.get('Answer'):
                answer_content_para = doc.add_paragraph(qa['Answer'])
                answer_content_para.paragraph_format.left_indent = Inches(0.5)
                answer_content_para.paragraph_format.first_line_indent = Inches(0.2)
                answer_content_para.paragraph_format.line_spacing = 1.6
                answer_content_para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # 空行
    
    def _add_key_points(self, doc: DocxDocument, summary_data: Dict):
        """添加要点提炼"""
        assistance = summary_data.get('meeting_assistance')
        if not assistance:
            return
        
        self._add_section_title(doc, '四、要点提炼')
        
        # 关键词
        keywords = assistance.get('keywords', [])
        if keywords:
            keywords_para = doc.add_paragraph()
            keywords_para.paragraph_format.left_indent = Inches(0.2)
            keywords_para.paragraph_format.space_before = Pt(3)
            kw_label = keywords_para.add_run('关键词：')
            kw_label.font.size = Pt(12)
            kw_label.font.bold = True
            kw_label.font.color.rgb = RGBColor(102, 51, 153)
            kw_content = keywords_para.add_run('、'.join(keywords))
            kw_content.font.size = Pt(11)
            keywords_para.paragraph_format.space_after = Pt(6)
        
        # 关键句
        key_sentences = assistance.get('key_sentences', [])
        if key_sentences:
            key_sentences_label = doc.add_paragraph()
            key_sentences_label.paragraph_format.left_indent = Inches(0.2)
            key_sentences_label.paragraph_format.space_before = Pt(6)
            ks_label_run = key_sentences_label.add_run('关键句：')
            ks_label_run.font.size = Pt(12)
            ks_label_run.font.bold = True
            ks_label_run.font.color.rgb = RGBColor(102, 51, 153)
            
            for idx, sentence in enumerate(key_sentences, 1):
                sentence_text = sentence.get('Text', '') if isinstance(sentence, dict) else str(sentence)
                if sentence_text:
                    sentence_para = doc.add_paragraph(f"{idx}. {sentence_text}")
                    sentence_para.paragraph_format.left_indent = Inches(0.5)
                    sentence_para.paragraph_format.first_line_indent = Inches(-0.2)
                    sentence_para.paragraph_format.line_spacing = 1.6
                    sentence_para.paragraph_format.space_after = Pt(3)
        
        # 待办事项
        actions = assistance.get('actions', [])
        if actions:
            actions_label = doc.add_paragraph()
            actions_label.paragraph_format.left_indent = Inches(0.2)
            actions_label.paragraph_format.space_before = Pt(6)
            act_label_run = actions_label.add_run('待办事项：')
            act_label_run.font.size = Pt(12)
            act_label_run.font.bold = True
            act_label_run.font.color.rgb = RGBColor(204, 102, 0)
            
            for idx, action in enumerate(actions, 1):
                action_text = action.get('Text', '') if isinstance(action, dict) else str(action)
                if action_text:
                    action_para = doc.add_paragraph(f"{idx}. {action_text}")
                    action_para.paragraph_format.left_indent = Inches(0.5)
                    action_para.paragraph_format.first_line_indent = Inches(-0.2)
                    action_para.paragraph_format.line_spacing = 1.6
                    action_para.paragraph_format.space_after = Pt(3)
        
        doc.add_paragraph()  # 空行
    
    def _add_mind_map(self, doc: DocxDocument, summary_data: Dict):
        """添加思维导图"""
        mind_map = summary_data.get('mind_map_summary', [])
        if not mind_map:
            return
        
        self._add_section_title(doc, '五、思维导图')
        
        mindmap_para = doc.add_paragraph('（思维导图结构请参考系统界面）')
        mindmap_para.paragraph_format.left_indent = Inches(0.3)
        mindmap_para.runs[0].font.size = Pt(11)
        mindmap_para.runs[0].font.italic = True
        mindmap_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()  # 空行

