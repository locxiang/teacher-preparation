"""
文档处理服务
"""
import os
import uuid
from typing import Optional, Tuple
from werkzeug.utils import secure_filename
import logging
from database import db
from models.document import Document

logger = logging.getLogger(__name__)

# 尝试导入docx解析库
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，docx解析功能将不可用。请运行: pip install python-docx")

# 尝试导入DashScope SDK用于AI提取摘要
try:
    from dashscope import Application
    from http import HTTPStatus
    from config import Config
    DASHSCOPE_AVAILABLE = True
except ImportError:
    DASHSCOPE_AVAILABLE = False
    logger.warning("DashScope SDK 未安装，AI提取摘要功能将不可用")
except Exception as e:
    DASHSCOPE_AVAILABLE = False
    logger.warning(f"DashScope SDK 配置错误: {e}")

# 允许的文档格式（备课资料仅支持 docx）
ALLOWED_EXTENSIONS = {'docx'}

# 文档文件存储目录
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads', 'documents')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 文件类型映射
FILE_TYPE_MAP = {
    'pdf': 'pdf',
    'doc': 'doc',
    'docx': 'docx',
    'ppt': 'ppt',
    'pptx': 'pptx',
    'txt': 'txt',
    'md': 'txt',
    'xls': 'xls',
    'xlsx': 'xlsx',
}

# MIME类型映射
MIME_TYPE_MAP = {
    'pdf': 'application/pdf',
    'doc': 'application/msword',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'ppt': 'application/vnd.ms-powerpoint',
    'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'txt': 'text/plain',
    'md': 'text/markdown',
    'xls': 'application/vnd.ms-excel',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
}


class DocumentService:
    """文档处理服务类"""
    
    def __init__(self):
        self.upload_folder = UPLOAD_FOLDER
    
    def allowed_file(self, filename: str) -> bool:
        """检查文件格式是否允许"""
        logger.debug(f"[检查文件格式] filename: {filename}")
        if '.' not in filename:
            logger.debug(f"[检查文件格式] 文件名没有扩展名: {filename}")
            return False
        
        file_ext_parts = filename.rsplit('.', 1)
        if len(file_ext_parts) < 2:
            logger.debug(f"[检查文件格式] 无法提取扩展名: {filename}")
            return False
        
        ext = file_ext_parts[1].lower()
        is_allowed = ext in ALLOWED_EXTENSIONS
        logger.debug(f"[检查文件格式] ext: {ext}, allowed: {is_allowed}, allowed_extensions: {ALLOWED_EXTENSIONS}")
        return is_allowed
    
    def get_file_type(self, filename: str) -> str:
        """获取文件类型"""
        logger.debug(f"[获取文件类型] filename: {filename}")
        if '.' not in filename:
            logger.warning(f"[获取文件类型] 文件名没有扩展名: {filename}")
            return 'unknown'
        
        file_ext_parts = filename.rsplit('.', 1)
        if len(file_ext_parts) < 2:
            logger.warning(f"[获取文件类型] 无法提取扩展名: {filename}")
            return 'unknown'
        
        ext = file_ext_parts[1].lower()
        file_type = FILE_TYPE_MAP.get(ext, 'unknown')
        logger.debug(f"[获取文件类型] ext: {ext}, file_type: {file_type}")
        return file_type
    
    def get_mime_type(self, filename: str) -> str:
        """获取MIME类型"""
        logger.debug(f"[获取MIME类型] filename: {filename}")
        if '.' not in filename:
            logger.warning(f"[获取MIME类型] 文件名没有扩展名: {filename}")
            return 'application/octet-stream'
        
        file_ext_parts = filename.rsplit('.', 1)
        if len(file_ext_parts) < 2:
            logger.warning(f"[获取MIME类型] 无法提取扩展名: {filename}")
            return 'application/octet-stream'
        
        ext = file_ext_parts[1].lower()
        mime_type = MIME_TYPE_MAP.get(ext, 'application/octet-stream')
        logger.debug(f"[获取MIME类型] ext: {ext}, mime_type: {mime_type}")
        return mime_type
    
    def save_document_file(self, file, meeting_id: str) -> Tuple[str, str]:
        """
        保存文档文件
        
        Args:
            file: 上传的文件对象
            meeting_id: 会议ID
        
        Returns:
            (文件路径, 文件名) 元组
        """
        logger.debug(f"[保存文件] 开始处理 - meeting_id: {meeting_id}, filename: {file.filename if file else None}")
        
        if not file or not file.filename:
            logger.error("[保存文件] 文件或文件名为空")
            raise ValueError("文件不能为空")
        
        logger.debug(f"[保存文件] 检查文件格式 - filename: {file.filename}")
        if not self.allowed_file(file.filename):
            logger.warning(f"[保存文件] 文件格式不支持 - filename: {file.filename}")
            raise ValueError("不支持的文档格式。备课资料仅支持 DOCX 格式（Word 文档）")
        
        # 生成安全的文件名
        original_filename = file.filename
        logger.debug(f"[保存文件] 原始文件名: {original_filename}")
        
        # 先从原始文件名提取扩展名（避免 secure_filename 清理掉扩展名）
        if '.' not in original_filename:
            logger.error(f"[保存文件] 原始文件名没有扩展名 - original_filename: {original_filename}")
            raise ValueError(f"文件名缺少扩展名: {original_filename}")
        
        original_ext_parts = original_filename.rsplit('.', 1)
        if len(original_ext_parts) < 2:
            logger.error(f"[保存文件] 无法从原始文件名提取扩展名 - original_filename: {original_filename}")
            raise ValueError(f"无法提取文件扩展名: {original_filename}")
        
        file_ext = original_ext_parts[1].lower()
        logger.debug(f"[保存文件] 文件扩展名: {file_ext}")
        
        # 验证扩展名是否允许（双重检查）
        if file_ext not in ALLOWED_EXTENSIONS:
            logger.warning(f"[保存文件] 文件扩展名不支持 - file_ext: {file_ext}, allowed: {ALLOWED_EXTENSIONS}")
            raise ValueError("不支持的文档格式。备课资料仅支持 DOCX 格式（Word 文档）")
        
        # 生成唯一文件名：使用 UUID 确保唯一性，不依赖原始文件名
        # 这样即使原始文件名包含特殊字符（如中文引号），也能正常保存
        unique_filename = f"{meeting_id}_{uuid.uuid4().hex[:8]}.{file_ext}"
        logger.debug(f"[保存文件] 唯一文件名: {unique_filename}")
        
        # 创建会议目录
        meeting_folder = os.path.join(self.upload_folder, meeting_id)
        logger.debug(f"[保存文件] 会议目录: {meeting_folder}")
        os.makedirs(meeting_folder, exist_ok=True)
        
        # 保存文件
        file_path = os.path.join(meeting_folder, unique_filename)
        logger.debug(f"[保存文件] 保存路径: {file_path}")
        file.save(file_path)
        
        logger.info(f"[保存文件] 文件保存成功 - file_path: {file_path}, size: {os.path.getsize(file_path)} 字节")
        
        return file_path, unique_filename
    
    def create_document_record(
        self,
        meeting_id: str,
        filename: str,
        original_filename: str,
        file_path: str,
        file_size: int,
        user_id: int
    ) -> Document:
        """
        创建文档记录
        
        Args:
            meeting_id: 会议ID
            filename: 存储的文件名
            original_filename: 原始文件名
            file_path: 文件路径
            file_size: 文件大小（字节）
            user_id: 用户ID
        
        Returns:
            Document对象
        """
        logger.debug(f"[创建文档记录] meeting_id: {meeting_id}, filename: {filename}, original_filename: {original_filename}, file_size: {file_size}")
        
        try:
            file_type = self.get_file_type(original_filename)
            mime_type = self.get_mime_type(original_filename)
            logger.debug(f"[创建文档记录] file_type: {file_type}, mime_type: {mime_type}")
            
            document = Document(
                meeting_id=meeting_id,
                filename=filename,
                original_filename=original_filename,
                file_path=file_path,
                file_size=file_size,
                file_type=file_type,
                mime_type=mime_type,
                status='uploaded',
                user_id=user_id
            )
            
            logger.debug(f"[创建文档记录] 准备添加到数据库")
            db.session.add(document)
            db.session.commit()
            
            logger.info(f"[创建文档记录] 文档记录创建成功 - document_id: {document.id}, file_type: {file_type}")
            
            return document
        except Exception as e:
            logger.error(f"[创建文档记录] 创建失败 - meeting_id: {meeting_id}, filename: {filename}, 错误: {str(e)}", exc_info=True)
            db.session.rollback()
            raise
    
    def get_documents_by_meeting(self, meeting_id: str, user_id: Optional[int] = None) -> list:
        """
        获取会议的所有文档
        
        Args:
            meeting_id: 会议ID
            user_id: 用户ID（可选，用于权限检查）
        
        Returns:
            文档列表
        """
        query = Document.query.filter_by(meeting_id=meeting_id)
        if user_id:
            query = query.filter_by(user_id=user_id)
        
        documents = query.order_by(Document.created_at.desc()).all()
        # 列表接口默认不包含内容（避免返回大量数据），但包含摘要
        return [doc.to_dict(include_content=False) for doc in documents]
    
    def delete_document(self, document_id: int, user_id: int) -> bool:
        """
        删除文档
        
        Args:
            document_id: 文档ID
            user_id: 用户ID（用于权限检查）
        
        Returns:
            是否删除成功
        """
        document = Document.query.filter_by(
            id=document_id,
            user_id=user_id
        ).first()
        
        if not document:
            return False
        
        # 删除文件
        if os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
            except Exception as e:
                logger.warning(f"删除文件失败: {e}")
        
        # 删除数据库记录
        db.session.delete(document)
        db.session.commit()
        
        logger.info(f"文档已删除: {document_id}")
        
        return True
    
    def update_document_status(
        self,
        document_id: int,
        status: str,
        parse_progress: Optional[int] = None,
        parsed_content: Optional[str] = None,
        summary: Optional[str] = None,
        error_message: Optional[str] = None
    ) -> Optional[Document]:
        """
        更新文档状态
        
        Args:
            document_id: 文档ID
            status: 状态
            parse_progress: 解析进度（0-100）
            parsed_content: 解析后的内容
            summary: AI提取的摘要和关键点
            error_message: 错误信息
        
        Returns:
            Document对象
        """
        document = Document.query.get(document_id)
        if not document:
            return None
        
        document.status = status
        if parse_progress is not None:
            document.parse_progress = parse_progress
        if parsed_content is not None:
            document.parsed_content = parsed_content
        if summary is not None:
            document.summary = summary
        if error_message is not None:
            document.error_message = error_message
        
        db.session.commit()
        
        return document
    
    def parse_docx(self, file_path: str) -> str:
        """
        解析docx文件，提取文本内容
        
        Args:
            file_path: docx文件路径
        
        Returns:
            提取的文本内容
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx 未安装，无法解析docx文件")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # 提取所有段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"解析docx文件失败: {file_path}, 错误: {str(e)}")
            raise ValueError(f"解析docx文件失败: {str(e)}")
    
    def extract_summary_with_ai(self, content: str, subject: Optional[str] = None, grade: Optional[str] = None) -> str:
        """
        使用AI提取备课资料的摘要和关键点
        
        Args:
            content: 文档内容
            subject: 学科（可选，用于上下文）
            grade: 年级（可选，用于上下文）
        
        Returns:
            AI提取的摘要和关键点
        """
        if not DASHSCOPE_AVAILABLE:
            logger.warning("DashScope SDK 不可用，跳过AI提取摘要")
            return ""
        
        if not Config.DASHSCOPE_API_KEY or not Config.DASHSCOPE_APP_ID:
            logger.warning("DashScope配置不完整，跳过AI提取摘要")
            return ""
        
        try:
            # 构建提示词
            context_parts = []
            if subject:
                context_parts.append(f"学科：{subject}")
            if grade:
                context_parts.append(f"年级：{grade}")
            
            context_str = "\n".join(context_parts) if context_parts else ""
            
            prompt = f"""你是一个AI助手，一个专门用于辅助备课教学的机器人。你不是老师，你只是一个AI助手，虽然你非常擅长备课教学，但你是一个机器人。请从以下备课资料中**提取**（不是生成或总结）核心信息点，这些信息将用于辅助备课会议讨论。

【上下文信息】
{context_str if context_str else "（未提供学科和年级信息）"}

【备课资料内容】
{content[:8000]}

【提取要求】
1. **只提取文档中明确存在的信息**，不要添加文档中没有的内容
2. **不要进行总结或概括**，而是提取具体的、结构化的关键信息点
3. 识别文档的核心主题和主要讨论点
4. 提取对备课会议讨论有价值的信息，包括但不限于：
   - 教学流程、步骤、方法
   - 注意事项、要点提醒
   - 关键概念、知识点
   - 教学资源、材料
   - 常见问题、难点
   - 经验分享、建议
   - 其他对备课有用的信息

【输出格式】
请按照以下格式输出，如果某项信息在文档中未提及，可以省略该项：

【核心主题】
（文档的核心主题是什么）

【关键信息点】
1. （提取的关键信息点1）
2. （提取的关键信息点2）
...

【重要提醒/注意事项】
（如果有的话）

【其他相关信息】
（其他对备课有用的信息）

请确保提取的信息都是文档中明确存在的，不要自行补充或推理。"""
            
            # 调用DashScope API
            response = Application.call(
                api_key=Config.DASHSCOPE_API_KEY,
                app_id=Config.DASHSCOPE_APP_ID,
                prompt=prompt
            )
            
            if response.status_code == HTTPStatus.OK:
                if hasattr(response, 'output') and hasattr(response.output, 'text'):
                    summary = response.output.text
                elif hasattr(response, 'output') and isinstance(response.output, dict):
                    summary = response.output.get('text', '')
                else:
                    summary = str(response.output) if hasattr(response, 'output') else str(response)
                
                logger.info(f"AI提取摘要成功，长度: {len(summary)}")
                return summary
            else:
                error_msg = f"AI提取摘要失败: {response.message if hasattr(response, 'message') else '未知错误'} (状态码: {response.status_code})"
                logger.error(error_msg)
                return ""
        except Exception as e:
            logger.error(f"AI提取摘要时出错: {str(e)}", exc_info=True)
            return ""
    
    def parse_and_extract_document(self, document_id: int, meeting_id: str) -> Optional[Document]:
        """
        解析文档并提取摘要（异步处理）
        
        Args:
            document_id: 文档ID
            meeting_id: 会议ID
        
        Returns:
            Document对象
        """
        document = Document.query.get(document_id)
        if not document:
            return None
        
        # 只处理docx文件
        if document.file_type != 'docx':
            logger.info(f"文档 {document_id} 不是docx格式，跳过解析")
            return document
        
        try:
            # 更新状态为处理中
            self.update_document_status(document_id, 'processing', parse_progress=10)
            
            # 解析docx文件
            logger.info(f"开始解析docx文件: {document.file_path}")
            parsed_content = self.parse_docx(document.file_path)
            self.update_document_status(document_id, 'processing', parse_progress=50, parsed_content=parsed_content)
            
            # 获取会议信息（用于AI提取摘要的上下文）
            from models.meeting import Meeting
            meeting = Meeting.query.get(meeting_id)
            subject = meeting.subject if meeting else None
            grade = meeting.grade if meeting else None
            
            # 使用AI提取摘要
            logger.info(f"开始使用AI提取摘要: {document_id}")
            summary = self.extract_summary_with_ai(parsed_content, subject, grade)
            
            # 分别存储解析内容和摘要
            self.update_document_status(
                document_id,
                'completed',
                parse_progress=100,
                parsed_content=parsed_content,
                summary=summary if summary else None
            )
            
            logger.info(f"文档解析和摘要提取完成: {document_id}")
            return document
            
        except Exception as e:
            error_msg = f"解析文档失败: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self.update_document_status(
                document_id,
                'failed',
                parse_progress=0,
                error_message=error_msg
            )
            return document

